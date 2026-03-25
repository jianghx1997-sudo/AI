from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from update_topic_library_links import refresh_topic_library_links


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS_DIR = ROOT / "outputs"
CANONICAL_REVIEW_DOCX = "01_\u5ba1\u6838\u7a3f.docx"
DEFAULT_FONT = "Microsoft YaHei"

H_TOPIC_INFO = "\u4e13\u9898\u4fe1\u606f"
H_TITLE = "\u6807\u9898"
H_COVER = "\u5c01\u9762"
H_COVER_POINTS = "\u5c01\u9762\u5217\u70b9"
H_CONFIRM = "\u786e\u8ba4\u8bf4\u660e"

L_TOPIC = "\u4e13\u9898\uff1a"
L_STATUS = "\u72b6\u6001\uff1a"
L_NOTE = "\u8bf4\u660e\uff1a"
L_COVER_TITLE = "\u4e3b\u6807\u9898\uff1a"
L_COVER_SUBTITLE = "\u526f\u6807\u9898\uff1a"
L_CARD_PREFIX = "\u77e5\u8bc6\u5361 "
L_CARD_TITLE = "\u6807\u9898\uff1a"
L_CARD_MARK = "\u6807\u8bb0\uff1a"
L_CARD_BODY = "\u6b63\u6587"
L_HASHTAGS = "#\u8bdd\u9898"


def ensure_font(run) -> None:
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)


def set_font_color(run, hex_color: str) -> None:
    run.font.color.rgb = RGBColor.from_string(hex_color)


def add_heading(doc: Document, text: str, size: int, color: str = "000000") -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(6)
    run = p.add_run(text)
    ensure_font(run)
    run.bold = True
    run.font.size = Pt(size)
    set_font_color(run, color)


def add_body_paragraph(doc: Document, text: str, *, numbered: bool = False, number: int | None = None) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.35
    fmt.space_after = Pt(10)
    fmt.space_before = Pt(0)
    if numbered and number is not None:
        fmt.left_indent = Pt(18)
        fmt.first_line_indent = Pt(-18)
        text = f"{number}. {text}"
    run = p.add_run(text)
    ensure_font(run)
    run.font.size = Pt(12)


def normalize_cover_points(points: list[Any]) -> list[str]:
    out: list[str] = []
    for point in points:
        if isinstance(point, str):
            value = point.strip()
        elif isinstance(point, dict):
            mark = str(point.get("mark_raw", "")).strip()
            text = str(point.get("text", "")).strip()
            value = f"{mark} {text}".strip() if mark else text
        else:
            value = ""
        if value:
            out.append(value)
    return out


def normalize_cards(cards: list[Any]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for idx, card in enumerate(cards, start=1):
        if not isinstance(card, dict):
            continue
        paragraphs = [str(item).strip() for item in card.get("paragraphs", []) if str(item).strip()]
        out.append(
            {
                "id": str(card.get("id", f"{idx:02d}")).strip() or f"{idx:02d}",
                "title": str(card.get("title", "")).strip(),
                "mark_raw": str(card.get("mark_raw", "")).strip(),
                "mark_render": str(card.get("mark_render", "")).strip(),
                "paragraphs": paragraphs,
            }
        )
    return out


def write_review_docx(source: dict[str, Any]) -> Path:
    date_text = str(source.get("date", "")).strip()
    topic = str(source.get("topic", "")).strip()
    title = str(source.get("title", "")).strip()
    cover = source.get("cover", {}) or {}
    cover_title = str(cover.get("title", "")).strip()
    cover_subtitle = str(cover.get("subtitle", "")).strip()
    cover_points = normalize_cover_points(cover.get("points", []) or [])
    cards = normalize_cards(source.get("cards", []) or [])
    hashtags = [str(tag).strip() for tag in source.get("hashtags", []) if str(tag).strip()]

    topic_dir = OUTPUTS_DIR / f"{date_text}_{topic}"
    topic_dir.mkdir(parents=True, exist_ok=True)
    out_path = topic_dir / CANONICAL_REVIEW_DOCX

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = DEFAULT_FONT
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    normal_style.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(f"{topic}\u5ba1\u6838\u7a3f")
    ensure_font(run)
    run.bold = True
    run.font.size = Pt(20)

    add_heading(doc, H_TOPIC_INFO, 16)
    add_body_paragraph(doc, f"{L_TOPIC} {topic}")
    add_body_paragraph(doc, f"{L_STATUS} \u5f85\u4eba\u5de5\u786e\u8ba4")
    add_body_paragraph(doc, f"{L_NOTE} \u8fd9\u4efd\u672c\u5730\u6587\u7a3f\u662f\u5f53\u524d\u5ba1\u6838\u9636\u6bb5\u7684\u552f\u4e00 source of truth\u3002")

    add_heading(doc, H_TITLE, 16)
    add_body_paragraph(doc, title)

    add_heading(doc, H_COVER, 16)
    add_body_paragraph(doc, f"{L_COVER_TITLE} {cover_title}")
    add_body_paragraph(doc, f"{L_COVER_SUBTITLE} {cover_subtitle}")
    add_heading(doc, H_COVER_POINTS, 14, color="8C4B2D")
    for point in cover_points:
        add_body_paragraph(doc, f"\u2022 {point}")

    for idx, card in enumerate(cards, start=1):
        add_heading(doc, f"{L_CARD_PREFIX}{idx}", 16)
        add_body_paragraph(doc, f"{L_CARD_TITLE} {card['title']}")
        add_body_paragraph(doc, f"{L_CARD_MARK} raw={card['mark_raw']} / render={card['mark_render']}")
        add_heading(doc, L_CARD_BODY, 14, color="8C4B2D")
        for p_idx, paragraph in enumerate(card["paragraphs"], start=1):
            add_body_paragraph(doc, paragraph, numbered=True, number=p_idx)

    add_heading(doc, L_HASHTAGS, 16)
    for tag in hashtags:
        add_body_paragraph(doc, f"\u2022 {tag}")

    add_heading(doc, H_CONFIRM, 16)
    add_body_paragraph(
        doc,
        "\u8bf7\u76f4\u63a5\u4fee\u6539\u8fd9\u4efd\u672c\u5730\u6587\u7a3f\u3002\u53ea\u6709\u5728\u660e\u786e\u8bf4\u201c\u6587\u7a3f\u65e0\u8bef\u201d\u6216\u201c\u53ef\u4ee5\u51fa\u56fe\u201d\u540e\uff0c\u624d\u80fd\u4ece\u8fd9\u4efd\u6587\u6863\u8fdb\u5165\u51fa\u56fe\u9636\u6bb5\u3002",
    )

    doc.save(out_path)
    refresh_topic_library_links()
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_json", type=Path)
    args = parser.parse_args()

    source = json.loads(args.source_json.read_text(encoding="utf-8-sig"))
    out_path = write_review_docx(source)
    print(out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
