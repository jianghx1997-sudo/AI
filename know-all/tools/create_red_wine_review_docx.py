from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
TOPIC = "红酒"
TOPIC_DIR = OUTPUTS / "20260323_红酒"
DOCX_PATH = TOPIC_DIR / "01_审核稿.docx"


TITLE = "红酒这5件事别想当然🍷"
MAIN_TITLE = "很多人喝红酒，但没真搞懂红酒"
SUBTITLE = "从颜色到醒酒，这几件事最容易想当然"

COVER_POINTS = [
    "🍇 红酒颜色不是果肉给的",
    "🍷 挂杯不等于酒更好",
    "⏳ 醒酒不是越久越好",
    "📅 红酒也不是越老越值钱",
    "🔩 螺旋盖不等于便宜酒",
]

CARDS = [
    {
        "title": "颜色不是果肉给的 🍇",
        "mark": "raw=🍇 / render=葡",
        "paragraphs": [
            "很多人一看到红酒是红的，就会很自然地觉得：那肯定是因为葡萄里面本来就是红色的汁。这个想法特别顺，因为“红葡萄酿红酒”听上去就像颜色会直接从果肉里流出来，很多人第一次接触红酒时，几乎都会这么理解。",
            "但大多数酿红酒的葡萄，挤出来的汁其实没有你想得那么红。红酒真正的颜色，主要来自葡萄皮。发酵时如果让果皮和葡萄汁一起待上一段时间，颜色、单宁和一部分风味就会慢慢被带出来，所以最后酒液才会变深。浸皮时间越长，通常颜色也会更深，口感也会更有抓力。",
            "也就是说，红酒之所以是红酒，关键不只是“用了红葡萄”，更重要的是酿的时候有没有让皮和汁一起发酵。很多人以为颜色是果肉天生带的，其实决定性更强的是工艺。"
        ],
    },
    {
        "title": "挂杯不等于酒更好 🍷",
        "mark": "raw=🍷 / render=杯",
        "paragraphs": [
            "很多人一晃酒杯，看到杯壁上慢慢往下流的“酒腿”，就会下意识觉得：这酒不错，挂杯这么明显。这个判断会流行，是因为挂杯特别直观，肉眼一下就能看到，天然容易被当成“高级感”的证据。",
            "但挂杯这件事，和“是不是好酒”并没有那么直接。它更多和酒精、糖分、甘油这些因素有关。酒精度高一点，或者残糖多一点，挂杯往往就更明显。所以你看到它“挂得漂亮”，不等于它在风味层次、平衡感、复杂度上就一定更强。它只能说明酒液的物理表现是这样，不能代替真正的品尝。",
            "红酒好不好，最后还是要回到闻起来香不香、喝起来酸度和单宁平不平衡、收尾干不干净。挂杯可以看，但它更像一个附带信号，不是整杯酒质量的总判决。"
        ],
    },
    {
        "title": "醒酒不是越久越好 ⏳",
        "mark": "raw=⏳ / render=醒",
        "paragraphs": [
            "一说到红酒，很多人就会自动把“醒酒”理解成一个固定动作，仿佛越高级的酒，就越该多醒一会儿。于是有些人一开瓶先放上一两个小时，觉得这样才算讲究，好像时间给得越久，酒就一定会越开、越好喝。",
            "但醒酒的作用，本质上是让酒和空气接触，帮助一些比较闷的香气慢慢打开，也让某些年轻、收得比较紧的酒喝起来没那么硬。可如果酒本身就比较轻、比较老，或者香气已经很脆弱，醒太久反而容易把它醒散了，果香掉得快，整体会变空。不是所有红酒都需要被“打开”得那么彻底。",
            "所以醒酒更像“看状态处理”，不是机械照做。年轻、封闭、结构重一点的酒可以多给一点时间；轻盈型或者老年份酒，很多时候反而应该边喝边观察，而不是一上来就长时间暴露在空气里。"
        ],
    },
    {
        "title": "红酒也不是越老越值钱 📅",
        "mark": "raw=📅 / render=年",
        "paragraphs": [
            "很多人一看到红酒瓶上的年份，就会很自然地觉得：年份越久，这酒应该就越厉害。这个印象很容易成立，因为“老酒更贵”在很多消费场景里本来就很常见，所以大家很容易把这种判断直接套到所有红酒身上。",
            "但红酒真不是都适合久放。大部分日常消费型红酒，本来就是为了在相对年轻的时候喝掉，讲究的是果味新鲜、结构顺口，不是靠十几年陈年去慢慢熬出来。如果把这类酒放太久，它最先掉下去的往往就是原本最讨喜的果香，最后剩下“放了很久”，不一定剩下“更好喝”。",
            "真正适合陈年的，通常得有比较好的酸度、单宁、酒精和整体平衡感，能撑得住时间慢慢改它。所以红酒能不能久放，关键不是“是不是红酒”，而是这瓶酒本身有没有那个底子。年份可以参考，但不能单独拿来代替品质判断。"
        ],
    },
    {
        "title": "螺旋盖不等于便宜酒 🔩",
        "mark": "raw=🔩 / render=盖",
        "paragraphs": [
            "很多人看到螺旋盖，第一反应还是会觉得：这酒是不是不太行，连木塞都没有。因为在很多人的印象里，木塞更“传统”，也更像高档红酒该有的样子，所以螺旋盖很容易被自动归进“便宜”“随便喝喝”的那一档。",
            "但这几年这个印象已经越来越站不住了。螺旋盖最大的好处，是稳定。它能减少木塞污染，也更方便控制密封状态。尤其是一些主打新鲜果香、希望尽快饮用的酒，螺旋盖反而很合适，因为酒厂更在意的是把风味稳稳保住，而不是额外制造一种“传统高级感”。对日常饮用型红酒来说，它常常是实用选择，不是廉价替代。",
            "所以“木塞更高级、螺旋盖更低端”很多时候是包装认知，不完全是酒本身的质量差异。你看到的是瓶口形式，酒厂真正考虑的往往是风格、稳定性和适饮期。"
        ],
    },
]

HASHTAGS = [
    "#红酒知识",
    "#葡萄酒入门",
    "#喝酒日常",
    "#生活常识",
    "#涨知识",
    "#每日闲聊语料库",
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_bookmark(paragraph, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "0")
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "0")
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11.5)
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    pf = normal.paragraph_format
    pf.line_spacing = 1.45
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)

    for style_name, size, bold, color in (
        ("Heading 1", 18, True, RGBColor(0x22, 0x22, 0x22)),
        ("Heading 2", 14, True, RGBColor(0x5A, 0x2F, 0x22)),
        ("Heading 3", 12.5, True, RGBColor(0x44, 0x44, 0x44)),
    ):
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_label_paragraph(doc: Document, label: str, value: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    r1 = p.add_run(label)
    set_east_asia_font(r1, "微软雅黑")
    r1.bold = True
    if value:
        r2 = p.add_run(value)
        set_east_asia_font(r2, "微软雅黑")


def add_body_paragraph(doc: Document, text: str, number: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.55
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.85)
    p.paragraph_format.first_line_indent = Cm(-0.65)
    prefix = p.add_run(f"{number}. ")
    set_east_asia_font(prefix, "微软雅黑")
    prefix.bold = True
    r = p.add_run(text)
    set_east_asia_font(r, "微软雅黑")


def build_doc() -> Document:
    doc = Document()
    style_document(doc)

    p = doc.add_paragraph(style="Heading 1")
    add_bookmark(p, "topic_review_root")
    r = p.add_run("每日闲聊语料库｜红酒 专题审核稿")
    set_east_asia_font(r, "微软雅黑")

    doc.add_paragraph("", style="Normal")

    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run("专题信息")
    set_east_asia_font(r, "微软雅黑")
    add_label_paragraph(doc, "话题：", TOPIC)
    add_label_paragraph(doc, "状态：", "审核稿，用户可直接修改本地文档")
    add_label_paragraph(doc, "说明：", "只有在你明确确认这份文稿无误后，才会从这份本地文档读取内容，进入图片生成环节。")

    doc.add_paragraph("", style="Normal")

    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run("标题")
    set_east_asia_font(r, "微软雅黑")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(TITLE)
    set_east_asia_font(r, "微软雅黑")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run("封面")
    set_east_asia_font(r, "微软雅黑")
    add_label_paragraph(doc, "主标题：", MAIN_TITLE)
    add_label_paragraph(doc, "副标题：", SUBTITLE)
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run("封面列点")
    set_east_asia_font(r, "微软雅黑")
    for item in COVER_POINTS:
        bullet = doc.add_paragraph(style="Normal")
        bullet.paragraph_format.left_indent = Cm(0.8)
        bullet.paragraph_format.first_line_indent = Cm(-0.3)
        bullet.paragraph_format.line_spacing = 1.35
        bullet.paragraph_format.space_after = Pt(4)
        r = bullet.add_run("• " + item)
        set_east_asia_font(r, "微软雅黑")

    for idx, card in enumerate(CARDS, 1):
        p = doc.add_paragraph(style="Heading 2")
        r = p.add_run(f"知识卡 {idx}")
        set_east_asia_font(r, "微软雅黑")
        add_label_paragraph(doc, "标题：", card["title"])
        add_label_paragraph(doc, "标记：", card["mark"])
        p = doc.add_paragraph(style="Heading 3")
        r = p.add_run("正文分段")
        set_east_asia_font(r, "微软雅黑")
        for pidx, paragraph in enumerate(card["paragraphs"], 1):
            add_body_paragraph(doc, paragraph, pidx)

    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run("#话题")
    set_east_asia_font(r, "微软雅黑")
    for tag in HASHTAGS:
        tag_p = doc.add_paragraph(style="Normal")
        tag_p.paragraph_format.space_after = Pt(2)
        r = tag_p.add_run(tag)
        set_east_asia_font(r, "微软雅黑")

    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run("确认说明")
    set_east_asia_font(r, "微软雅黑")
    add_body_paragraph(
        doc,
        "你可以直接修改这份文档。等你明确告诉我“文稿无误”或“可以出图”之后，我会以这份本地文档为准读取内容，再进入图片生成阶段。",
        1,
    )

    return doc


def main() -> int:
    TOPIC_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(DOCX_PATH)
    print(DOCX_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
